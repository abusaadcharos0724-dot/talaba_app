from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_referat_docx(topic: str, content: str, output_path: str):
    doc = Document()
    
    # Title
    title = doc.add_heading(topic, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content
    # We assume content might have sections separated by newlines or markdown-like headers
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('###') or line.startswith('**'):
            p = doc.add_heading(line.replace('#', '').replace('*', '').strip(), level=2)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.runs[0] if p.runs else p.add_run()
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    doc.save(output_path)
    return output_path
